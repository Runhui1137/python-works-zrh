import docx
import os
from win32com import client
import csv
import re

# 配置信息
docx_path = "D:\\2020春课堂\\DOTNET\\实验1\\软件182-实验一"        # 目标文件夹的路径
sno_patt = u'2018\d\d\d\d\d\d'                                 # 正则表达式, 用于在docx中查找学号
sname_patt = u'姓名\s*\S*'                                      # 正则表达式, 用于在docx中查找姓名
# 要格式化实验报告文件名, 直接修改上免变量的值即可
#########################################################################
word = client.Dispatch("Word.Application")    # 用于doc格式转换为docx的工具

def get_files():
    return os.listdir(docx_path)

# modify_single_file: 对单个文件进行改名
## filename: 文件名(不含路径)
def modify_single_file(filename):
    # 判断并记录该文件的类型
    is_docx = True
    if (filename.split(".")[-1] == 'doc'):
        is_docx = False
    elif (filename.split(".")[-1] == 'docx'):
        is_docx = True
    else:
        return
    # 得到该文件的路径
    if docx_path[-1] != '\\' :
        filepath = docx_path + '\\'
    filepath += filename
    docx_filepath = filepath     # 变量docx_filepath是为了解决doc文件转换而存在, 如果该文件为docx格式, 该变量与filepath值相同
    extension = 'docx'          # 变量extension即为文件之前的扩展名(doc或docx)
    # 如果不为docx文件, 则生成临时docx文件
    if not is_docx :
        docx_filepath = save_single_doc_as_docx(filepath, filename)
        extension = 'doc'
    # 打开并遍历docx, 使用正则表达式进行匹配学号
    re_ans = None
    dx = docx.Document(docx=docx_filepath)
    sno = ""
    sname = ""
    for line in dx.paragraphs :
        re_ans = re.search(string=line.text, pattern=sno_patt)
        if re_ans != None:
            sno = re_ans.group().split(" ")[-1]
            break

    for line in dx.paragraphs :
        re_ans = re.search(string=line.text, pattern=sname_patt)
        if re_ans != None:
            sname = re_ans.group().split(" ")[-1]
            break

    new_name = ''
    if re_ans != None:
        ###### 在这里可以修改重命名文件的格式
        new_name = "{0} {1}.{2}".format(sname, sno, extension)
        new_path = docx_path+ "\\" + new_name
        os.rename(filepath, new_path)
    # 删除临时生成的docx文件
    if not is_docx:
        os.remove(docx_filepath)
    print("{0}    --->    {1}      success !".format(filename, new_name))

# 将doc文件另存为docx文件至temp目录中
def save_single_doc_as_docx(filepath, filename):
    docx_filepath = docx_path + '\\temp\\' + filename.split('.')[0] + '.docx'
    doc = word.Documents.Open(filepath)
    doc.SaveAs(docx_filepath, 12)
    doc.Close()
    return docx_filepath

def main():
    global docx_path
    docx_path = os.path.abspath(docx_path)
    docx_path = "\\".join(docx_path.split('/'))
    if not os.path.exists(docx_path + "\\temp"):
        os.mkdir(docx_path + "\\temp")
    # 获取目标文件夹下的所有文件
    print(docx_path)
    file_list = get_files()
    # 对每个文件重命名
    for file in file_list:
        modify_single_file(file)
    print("All Tasks Success!")
    # 析构win32com
    # 在没使用完win32com模块前千万别调用Quit方法, 否则会出现错误
    word.Quit()

main()
